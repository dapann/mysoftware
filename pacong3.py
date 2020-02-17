# -*- coding:utf-8 -*-
##******** 字典多标题，多页读取 **********
#
import urllib
import urllib.request
import os
import random
import re
from bs4 import BeautifulSoup
import time

from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from  docx.shared import Inches
from  docx.shared import RGBColor

def color(value):
	digit = list(map(str, range(10))) + list("ABCDEF")
	if isinstance(value, tuple):
		string = '#'
		for i in value:
			a1 = i // 16
			a2 = i % 16
			string += digit[a1] + digit[a2]
			return string
	elif isinstance(value, str):
		a1 = digit.index(value[1]) * 16 + digit.index(value[2])
		a2 = digit.index(value[3]) * 16 + digit.index(value[4])
		a3 = digit.index(value[5]) * 16 + digit.index(value[6])
		return [a1, a2, a3]
def colour(str):
	dict={"pink":[255,192,203],"crimson":[220,20,60],"lavenderblush":[255,240,245],"palevioletred":[219,112,147],"hotpink":[255,105,180],"mediumvioletred":[199,21,133],"orchid":[218,112,214],"thistle":[216,191,216],"plum":[221,160,221],"violet":[238,130,238],"magenta":[255,0,255],"fuchsia":[255,0,255],"darkmagenta":[139,0,139],"purple":[128,0,128],"mediumorchid":[186,85,211],"darkviolet":[148,0,211],"indigo":[75,0,130],"blueviolet":[138,43,226],"mediumpurple":[147,112,219],"mediumslateblue":[123,104,238],"slateblue":[106,90,205],"darkslateblue":[72,61,139],"lavender":[230,230,250],"ghostwhite":[248,248,255],"blue":[0,0,255],"mediumblue":[0,0,205],"midnightblue":[25,25,112],"darkblue":[0,0,139],"navy":[0,0,128],"royalblue":[65,105,225],"cornflowerblue":[100,149,237],"lightsteelblue":[176,196,222],"lightslategray":[119,136,153],"slategray":[112,128,144],"dodgerblue":[30,144,255],"aliceblue":[240,248,255],"steelblue":[70,130,180],"lightskyblue":[135,206,250],"skyblue":[135,206,235],"deepskyblue":[0,191,255],"lightblue":[173,216,230],"powderblue":[176,224,230],"cadetblue":[95,158,160],"azure":[240,255,255],"lightcyan":[224,255,255],"paleturquoise":[175,238,238],"cyan":[0,255,255],"aqua":[0,255,255],"darkturquoise":[0,206,209],"darkslategray":[47,79,79],"darkcyan":[0,139,139],"teal":[0,128,128],"mediumturquoise":[72,209,204],"lightseagreen":[32,178,170],"turquoise":[64,224,208],"aquamarine":[127,255,212],"mediumaquamarine":[102,205,170],"mediumspringgreen":[0,250,154],"mintcream":[245,255,250],"springgreen":[0,255,127],"mediumseagreen":[60,179,113],"seagreen":[46,139,87],"honeydew":[240,255,240],"lightgreen":[144,238,144],"palegreen":[152,251,152],"darkseagreen":[143,188,143],"limegreen":[50,205,50],"lime":[0,255,0],"forestgreen":[34,139,34],"chartreuse":[127,255,0],"lawngreen":[124,252,0],"greenyellow":[173,255,47],"darkolivegreen":[85,107,47],"yellowgreen":[154,205,50],"olivedrab":[107,142,35],"beige":[245,245,220],"lightgoldenrodyellow":[250,250,210],"ivory":[255,255,240],"lightyellow":[255,255,224],"yellow":[255,255,0],"olive":[128,128,0],"darkkhaki":[189,183,107],"lemonchiffon":[255,250,205],"palegoldenrod":[238,232,170],"khaki":[240,230,140],"gold":[255,215,0],"cornsilk":[255,248,220],"goldenrod":[218,165,32],"darkgoldenrod":[184,134,11],"floralwhite":[255,250,240],"oldlace":[253,245,230],"wheat":[245,222,179],"moccasin":[255,228,181],"orange":[255,165,0],"papayawhip":[255,239,213],"blanchedalmond":[255,235,205],"navajowhite":[255,222,173],"antiquewhite":[250,235,215],"tan":[210,180,140],"burlywood":[222,184,135],"bisque":[255,228,196],"darkorange":[255,140,0],"linen":[250,240,230],"peru":[205,133,63],"sandybrown":[244,164,96],"chocolate":[210,105,30],"chocolatesaddlebrown":[192,14,235],"seashell":[255,245,238],"sienna":[160,82,45],"lightsalmon":[255,160,122],"coral":[255,127,80],"orangered":[255,69,0],"tomato":[255,99,71],"mistyrose":[255,228,225],"salmon":[250,128,114],"snow":[255,250,250],"lightcoral":[240,128,128],"rosybrown":[188,143,143],"indianred":[205,92,92],"red":[255,0,0],"brown":[165,42,42],"firebrick":[178,34,34],"darkred":[139,0,0],"maroon":[128,0,0],"white":[255,255,255],"whitesmoke":[245,245,245],"gainsboro":[220,220,220],"lightgrey":[211,211,211],"silver":[192,192,192],"darkgray":[169,169,169],"dimgray":[105,105,105],"black":[0,0,0]}

	result=dict.get(str,(0,0,0))
	return result

#打开文档
document = Document()
bold=''
txtfont=[]

md={
'218942':2,
	}

for k in sorted(md.keys()):
	for i in range(1,md[k]+1):
		url="http://www.ngotcmszh.com/thread-"+k+"-"+str(i)+"-1.html" 		# 要爬取的网络地址
		my_headers = [
					"Mozilla/5.0 (Windows NT6.1;WOW64;rv:27.0)Gecko/20100101 Firefox/27.0",
					"Mozilla/5.0 (Windows;U;Windows NT 6.1;en-US;rv:1.9.1.6)Gecko/20091201 Firefox/3.5.6",
					"Mozilla/5.0 (Windows;U;Windows NT 6.1)AppleWebKit/537.36(KHTML,like Gecko) Chrome/34.0.1838.2 Safari/",
					"Mozilla/5.0 (X11;Ubuntu; Linux i686;rv:10.0) Gecko/20100101 Firefox/10.0",
					"Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; AcooBrowser; .NET CLR 1.1.4322; .NET CLR 2.0.50727)",
				    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 6.0; Acoo Browser; SLCC1; .NET CLR 2.0.50727; Media Center PC 5.0; .NET CLR 3.0.04506)",
				    "Mozilla/4.0 (compatible; MSIE 7.0; AOL 9.5; AOLBuild 4337.35; Windows NT 5.1; .NET CLR 1.1.4322; .NET CLR 2.0.50727)",
				    "Mozilla/5.0 (Windows; U; MSIE 9.0; Windows NT 9.0; en-US)",
				    "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; Win64; x64; Trident/5.0; .NET CLR 3.5.30729; .NET CLR 3.0.30729; .NET CLR 2.0.50727; Media Center PC 6.0)",
				    "Mozilla/5.0 (compatible; MSIE 8.0; Windows NT 6.0; Trident/4.0; WOW64; Trident/4.0; SLCC2; .NET CLR 2.0.50727; .NET CLR 3.5.30729; .NET CLR 3.0.30729; .NET CLR 1.0.3705; .NET CLR 1.1.4322)",
				    "Mozilla/4.0 (compatible; MSIE 7.0b; Windows NT 5.2; .NET CLR 1.1.4322; .NET CLR 2.0.50727; InfoPath.2; .NET CLR 3.0.04506.30)",
				    "Mozilla/5.0 (Windows; U; Windows NT 5.1; zh-CN) AppleWebKit/523.15 (KHTML, like Gecko, Safari/419.3) Arora/0.3 (Change: 287 c9dfb30)",
				    "Mozilla/5.0 (X11; U; Linux; en-US) AppleWebKit/527+ (KHTML, like Gecko, Safari/419.3) Arora/0.6",
				    "Mozilla/5.0 (Windows; U; Windows NT 5.1; en-US; rv:1.8.1.2pre) Gecko/20070215 K-Ninja/2.1.1",
				    "Mozilla/5.0 (Windows; U; Windows NT 5.1; zh-CN; rv:1.9) Gecko/20080705 Firefox/3.0 Kapiko/3.0",
				    "Mozilla/5.0 (X11; Linux i686; U;) Gecko/20070322 Kazehakase/0.4.5",
				    "Mozilla/5.0 (X11; U; Linux i686; en-US; rv:1.9.0.8) Gecko Fedora/1.9.0.8-1.fc10 Kazehakase/0.5.6",
				    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.56 Safari/535.11",
				    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_7_3) AppleWebKit/535.20 (KHTML, like Gecko) Chrome/19.0.1036.7 Safari/535.20",
				    "Opera/9.80 (Macintosh; Intel Mac OS X 10.6.8; U; fr) Presto/2.9.168 Version/11.52",
		]

		random_header = random.choice(my_headers)

		req = urllib.request.Request(url)
		req.add_header("User-Agent",random_header)
		req.add_header("Host","www.ngotcmszh.com")
		req.add_header("GET",url)

		htmq = urllib.request.urlopen(req)
		print(htmq.status) 							#读取网页响应码，200表示正确响应

		htmx = htmq.read().decode('gb18030','ignore')
		# html=re.sub(r"<font color=\"(?:red|#ED220B|#c00000)\">(.+?)</font>","{"+"\g<1>"+"}",htmx)

		soup=BeautifulSoup(htmx,'lxml')  # 使用html解析器进行解析

		title="标题/"+k+"/："+soup.title.string.split('-')[0]
		print(title)

		if i==1:
			document.add_heading(title,1)

		contents=soup.find_all("div", class_=["t_fsz","authi","cm"])
		for content in contents:
			if content['class']==['authi']:
				if 'xw1' in str(content.a):
					# print "["+content.a.get_text()+"]"
					#添加文本
					document.add_paragraph(u"["+content.a.get_text()+"]")
			else:
				text = str(content)
				pattern = re.compile(r'>[^<]+<|<strong>|<font color=[^>]+>')
				pos = 0
				while True:
					match = pattern.search(text, pos)
					if not match:
						break
					s = match.start()
					e = match.end()

					if text[s:e][1]=='s':
						bold='b'
					elif text[s:e][1]=='f':
						fontcode=text[s:e].replace('<font color="','').replace('">','')

						if fontcode[0]=='#':
							txtfont=color(fontcode.upper().ljust(7,'0'))
						else:
							txtfont=colour(fontcode)
					else:
						lines=text[s:e].replace('<','').replace('>','')
						paragraph =document.add_paragraph().add_run(lines)

						font = paragraph.font
						if bold=='b':
							font.bold=True
						if txtfont!=[]:
							txtcolor = font.color
							txtcolor.rgb = RGBColor(txtfont[0],txtfont[1],txtfont[2])
						bold=''
						txtfont=[]

					# Move forward in text for the next search
					pos = e-1
	time.sleep(10)
document.save('demo.docx')

print("Done!")
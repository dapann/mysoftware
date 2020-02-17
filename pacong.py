# -*- coding:utf-8 -*-  
import urllib
import os
import random
import re 
from bs4 import BeautifulSoup

from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from  docx.shared import Inches

from docx.shared import RGBColor  

import sys
reload(sys)
sys.setdefaultencoding('utf-8')

document = Document()	

uid="184715"

for i in range(1,3):
	url="http://www.ngotcmszh.com/thread-"+uid+"-"+str(i)+"-1.html" 		# 要爬取的网络地址


	htmx=urllib.urlopen(url).read()											# 将网页源代码赋予menuCode

	html=re.sub(r"<font color=\"(?:red|#ED220B|#c00000)\">(.+?)</font>","{"+"\g<1>"+"}",htmx)


	soup=BeautifulSoup(html,'lxml')  # 使用html解析器进行解析

	title="标题："+soup.title.string.split('-')[0]
	print "标题："+soup.title.string.split('-')[0]
	
	if i==1:
		document.add_heading(title,1)

	contents=soup.find_all("div", class_=["t_fsz","authi","cm"]) 
	for content in contents:
	# print content.get_text()
		if content['class']==['authi']:
			if 'xw1' in str(content.a):
				# print "["+content.a.get_text()+"]"
				#添加文本
				document.add_paragraph(u"["+content.a.get_text()+"]")
		else:
			body=content.get_text().replace("}{","")

			for b in body.split("{"):
				if '}' in b:
					line1=b.split("}")[0]
					line2=b.split("}")[1]

					paragraph =document.add_paragraph().add_run(line1)
					font = paragraph.font  
					font.color.rgb = RGBColor(0xff, 0x00 , 0x00)

					document.add_paragraph(line2)
				else:
					document.add_paragraph(b)

document.save('demo.docx')

print "Done!"
    	




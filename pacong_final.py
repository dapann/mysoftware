# -*- coding:utf-8 -*- 
# 读取本地html，提取目标位置及页数
# 

import urllib
import urllib2
import os
import random
import re 
from bs4 import BeautifulSoup

from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from  docx.shared import Inches

import sys
reload(sys)
sys.setdefaultencoding('utf-8')

document = Document()	

for a in range(5,0,-1):
	path="E:/testing/testing-"+str(a).rjust(3,'0')+".html"
	d={}
	html=open(path).read()

	soup=BeautifulSoup(html,'lxml')
	contents=soup.find_all("th") #"span",class_="tps"
	for content in contents:
		if content.a!=None:
			key=content.a['href'].split('-')[1]
			if content.span!=None:
				if str(content.span.get_text()).strip()[-1]!=None and content.span.get_text()!="已关闭":
					ment=int(str(content.span.get_text()).strip()[-1])
					d[key]=ment

			else:
				d[key]=1
	for key in sorted(d,reverse=False):
		print key+"-"+str(d[key])
		for i in range(1,d[key]+1):
			url="http://www.ngotcm.com/forum/thread-"+key+"-"+str(i)+"-1.html" 		# 要爬取的网络地址
			my_headers = [  
						"Mozilla/5.0 (Windows NT6.1;WOW64;rv:27.0)Gecko/20100101 Firefox/27.0",  
						"Mozilla/5.0 (Windows;U;Windows NT 6.1;en-US;rv:1.9.1.6)Gecko/20091201 Firefox/3.5.6",  
						"Mozilla/5.0 (Windows;U;Windows NT 6.1)AppleWebKit/537.36(KHTML,like Gecko) Chrome/34.0.1838.2 Safari/",  
						"Mozilla/5.0 (X11;Ubuntu; Linux i686;rv:10.0) Gecko/20100101 Firefox/10.0",
						"Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; AcooBrowser; .NET CLR 1.1.4322; .NET CLR 2.0.50727)",
					    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 6.0; Acoo Browser; SLCC1; .NET CLR 2.0.50727; Media Center PC 5.0; .NET CLR 3.0.04506)",
					    "Mozilla/4.0 (compatible; MSIE 7.0; AOL 9.5; AOLBuild 4337.35; Windows NT 5.1; .NET CLR 1.1.4322; .NET CLR 2.0.50727)",
					    "Mozilla/5.0 (Windows; U; MSIE 9.0; Windows NT 9.0; en-US)",
					    "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; Win64; x64; Trident/5.0; .NET CLR 3.5.30729; .NET CLR 3.0.30729; .NET CLR 2.0.50727; Media Center PC 6.0)",
					    "Mozilla/5.0 (compatible; MSIE 8.0; Windows NT 6.0; Trident/4.0; WOW64; Trident/4.0; SLCC2; .NET CLR 2.0.50727; .NET CLR 3.5.30729; .NET CLR 3.0.30729; .NET CLR 1.0.3705; .NET CLR 1.1.4322)",
					    "Mozilla/4.0 (compatible; MSIE 7.0b; Windows NT 5.2; .NET CLR 1.1.4322; .NET CLR 2.0.50727; InfoPath.2; .NET CLR 3.0.04506.30)",
					    "Mozilla/5.0 (Windows; U; Windows NT 5.1; zh-CN) AppleWebKit/523.15 (KHTML, like Gecko, Safari/419.3) Arora/0.3 (Change: 287 c9dfb30)",
					    "Mozilla/5.0 (X11; U; Linux; en-US) AppleWebKit/527+ (KHTML, like Gecko, Safari/419.3) Arora/0.6",
					    "Mozilla/5.0 (Windows; U; Windows NT 5.1; en-US; rv:1.8.1.2pre) Gecko/20070215 K-Ninja/2.1.1",
					    "Mozilla/5.0 (Windows; U; Windows NT 5.1; zh-CN; rv:1.9) Gecko/20080705 Firefox/3.0 Kapiko/3.0",
					    "Mozilla/5.0 (X11; Linux i686; U;) Gecko/20070322 Kazehakase/0.4.5",
					    "Mozilla/5.0 (X11; U; Linux i686; en-US; rv:1.9.0.8) Gecko Fedora/1.9.0.8-1.fc10 Kazehakase/0.5.6",
					    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.56 Safari/535.11",
					    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_7_3) AppleWebKit/535.20 (KHTML, like Gecko) Chrome/19.0.1036.7 Safari/535.20",
					    "Opera/9.80 (Macintosh; Intel Mac OS X 10.6.8; U; fr) Presto/2.9.168 Version/11.52",
]  
 
			random_header = random.choice(my_headers)  
			      
			req = urllib2.Request(url)  
			req.add_header("User-Agent",random_header)  
			req.add_header("Host","www.ngotcm.com")  
			req.add_header("GET",url)  
			      
			htmx = urllib2.urlopen(req).read()

			html=re.sub(r"<font color=\"(?:red|#ED220B|#c00000)\">(.+?)</font>","{"+"\g<1>"+"}",htmx)

			soup=BeautifulSoup(html,'lxml')  # 使用html解析器进行解析

			title="标题："+soup.title.string.split('-')[0]
			print title
		
			if i==1:
				document.add_heading(title,1)

			contents=soup.find_all("div", class_=["t_fsz","authi","cm"]) 
			for content in contents:
				if content['class']==['authi']:
					if 'xw1' in str(content.a):
						line1= "["+content.a.get_text()+"]"
					#添加文本
						paragraph = document.add_paragraph(line1)
				else:
					line2=content.get_text()
					paragraph = document.add_paragraph(line2)

document.save('demo.docx')

print "Done!"
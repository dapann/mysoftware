# -*- coding:utf-8 -*-  
#******** 单页读取 *********
import urllib.request
import re 
from bs4 import BeautifulSoup

from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from  docx.shared import Inches
from  docx.shared import RGBColor

def color(value):    
	digit = list(map(str, range(10))) + list("ABCDEF")
	a1 = digit.index(value[1]) * 16 + digit.index(value[2])        
	a2 = digit.index(value[3]) * 16 + digit.index(value[4])        
	a3 = digit.index(value[5]) * 16 + digit.index(value[6])
	return [a1, a2, a3]
def colour(str):
	dict={"pink":[255,192,203],"crimson":[220,20,60],"lavenderblush":[255,240,245],"palevioletred":[219,112,147],"hotpink":[255,105,180],"mediumvioletred":[199,21,133],"orchid":[218,112,214],"thistle":[216,191,216],"plum":[221,160,221],"violet":[238,130,238],"magenta":[255,0,255],"fuchsia":[255,0,255],"darkmagenta":[139,0,139],"purple":[128,0,128],"mediumorchid":[186,85,211],"darkviolet":[148,0,211],"indigo":[75,0,130],"blueviolet":[138,43,226],"mediumpurple":[147,112,219],"mediumslateblue":[123,104,238],"slateblue":[106,90,205],"darkslateblue":[72,61,139],"lavender":[230,230,250],"ghostwhite":[248,248,255],"blue":[0,0,255],"mediumblue":[0,0,205],"midnightblue":[25,25,112],"darkblue":[0,0,139],"navy":[0,0,128],"royalblue":[65,105,225],"cornflowerblue":[100,149,237],"lightsteelblue":[176,196,222],"lightslategray":[119,136,153],"slategray":[112,128,144],"dodgerblue":[30,144,255],"aliceblue":[240,248,255],"steelblue":[70,130,180],"lightskyblue":[135,206,250],"skyblue":[135,206,235],"deepskyblue":[0,191,255],"lightblue":[173,216,230],"powderblue":[176,224,230],"cadetblue":[95,158,160],"azure":[240,255,255],"lightcyan":[224,255,255],"paleturquoise":[175,238,238],"cyan":[0,255,255],"aqua":[0,255,255],"darkturquoise":[0,206,209],"darkslategray":[47,79,79],"darkcyan":[0,139,139],"teal":[0,128,128],"mediumturquoise":[72,209,204],"lightseagreen":[32,178,170],"turquoise":[64,224,208],"aquamarine":[127,255,212],"mediumaquamarine":[102,205,170],"mediumspringgreen":[0,250,154],"mintcream":[245,255,250],"springgreen":[0,255,127],"mediumseagreen":[60,179,113],"seagreen":[46,139,87],"honeydew":[240,255,240],"lightgreen":[144,238,144],"palegreen":[152,251,152],"darkseagreen":[143,188,143],"limegreen":[50,205,50],"lime":[0,255,0],"forestgreen":[34,139,34],"chartreuse":[127,255,0],"lawngreen":[124,252,0],"greenyellow":[173,255,47],"darkolivegreen":[85,107,47],"yellowgreen":[154,205,50],"olivedrab":[107,142,35],"beige":[245,245,220],"lightgoldenrodyellow":[250,250,210],"ivory":[255,255,240],"lightyellow":[255,255,224],"yellow":[255,255,0],"olive":[128,128,0],"darkkhaki":[189,183,107],"lemonchiffon":[255,250,205],"palegoldenrod":[238,232,170],"khaki":[240,230,140],"gold":[255,215,0],"cornsilk":[255,248,220],"goldenrod":[218,165,32],"darkgoldenrod":[184,134,11],"floralwhite":[255,250,240],"oldlace":[253,245,230],"wheat":[245,222,179],"moccasin":[255,228,181],"orange":[255,165,0],"papayawhip":[255,239,213],"blanchedalmond":[255,235,205],"navajowhite":[255,222,173],"antiquewhite":[250,235,215],"tan":[210,180,140],"burlywood":[222,184,135],"bisque":[255,228,196],"darkorange":[255,140,0],"linen":[250,240,230],"peru":[205,133,63],"sandybrown":[244,164,96],"chocolate":[210,105,30],"chocolatesaddlebrown":[192,14,235],"seashell":[255,245,238],"sienna":[160,82,45],"lightsalmon":[255,160,122],"coral":[255,127,80],"orangered":[255,69,0],"tomato":[255,99,71],"mistyrose":[255,228,225],"salmon":[250,128,114],"snow":[255,250,250],"lightcoral":[240,128,128],"rosybrown":[188,143,143],"indianred":[205,92,92],"red":[255,0,0],"brown":[165,42,42],"firebrick":[178,34,34],"darkred":[139,0,0],"maroon":[128,0,0],"white":[255,255,255],"whitesmoke":[245,245,245],"gainsboro":[220,220,220],"lightgrey":[211,211,211],"silver":[192,192,192],"darkgray":[169,169,169],"dimgray":[105,105,105],"black":[0,0,0]}
	
	result=dict.get(str,(0,0,0))
	return result

#打开文档
document = Document()
bold=''
txtfont=[]		

url="http://www.ngotcmszh.com/thread-216924-3-1.html"
htmx=urllib.request.urlopen(url).read().decode('GBK')									# 将网页源代码赋予menuCode

# htmx=open(path).read()
# html=re.sub(r"<font color=\"(?:red|#ED220B|#c00000)\">(.+?)</font>","{"+"\g<1>"+"}",htmx)


soup=BeautifulSoup(htmx,'lxml')  # 使用html解析器进行解析

contents=soup.find_all("div", class_=["t_fsz","authi","cm"]) 
for content in contents:
	if content['class']==['authi']:
		if 'xw1' in str(content.a):
			# print "["+content.a.get_text()+"]"
			#添加文本
			document.add_paragraph(u"["+content.a.get_text()+"]")
	else:
		text = str(content)
		pattern = re.compile(r'>[^<]+<|<strong>|<font color=[^>]+>') 
		pos = 0
		while True:    
			match = pattern.search(text, pos)    
			if not match:        
				break    
			s = match.start()    
			e = match.end()    

			if text[s:e][1]=='s':
				bold='b'
			elif text[s:e][1]=='f':
				fontcode=text[s:e].replace('<font color="','').replace('">','')

				if fontcode[0]=='#':
					txtfont=color(fontcode.upper().ljust(7,'0'))
				else:
					txtfont=colour(fontcode)
			else:
				lines=text[s:e].replace('<','').replace('>','')
				paragraph =document.add_paragraph().add_run(lines)
				
				font = paragraph.font 
				if bold=='b':
					font.bold=True
				if txtfont!=[]:
					txtcolor = font.color
					txtcolor.rgb = RGBColor(txtfont[0],txtfont[1],txtfont[2])
				bold=''
				txtfont=[]

			# Move forward in text for the next search
			pos = e-1
#保存文件
document.save('demo_singlepage.docx')

print("Done!")